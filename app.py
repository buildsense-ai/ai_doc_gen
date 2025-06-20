#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FastAPI Web应用：AI文档生成器
集成聊天界面和文档生成功能
"""

import os
import json
import logging
import uuid
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Body
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from starlette.requests import Request
from pydantic import BaseModel
from docx import Document

# 导入现有的AI文档生成器
from main import AIDocGenerator

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# --- FastAPI App Setup ---
app = FastAPI(title="AI文档生成器", description="智能文档生成平台")
app.mount("/static", StaticFiles(directory="frontend/static"), name="static")
templates = Jinja2Templates(directory="frontend/templates")

# --- Data Models ---
class DocumentItem:
    def __init__(self, name: str, doc_id: str):
        self.id = doc_id
        self.name = name
        self.status = "待处理"  # "待处理", "已完成", "生成失败"
        self.filled_document_path: Optional[str] = None
        self.error_info: Optional[str] = None
        self.matched_template_path: Optional[str] = None

class ChatSession:
    def __init__(self, session_id: str):
        self.session_id = session_id
        # Key: doc_id, Value: DocumentItem object
        self.document_items: Dict[str, DocumentItem] = {}
        # Key: template_name, Value: template_path
        self.templates: Dict[str, str] = {}
        # Key: file_name, Value: file_path
        self.context_files: Dict[str, str] = {}
    
    def get_dashboard_data(self) -> List[Dict]:
        """获取用于前端展示的文档清单数据"""
        items_data = []
        for item in self.document_items.values():
            item_data = item.__dict__.copy()
            item_data['template_status'] = "已匹配" if item.matched_template_path else "未匹配"
            items_data.append(item_data)
        return sorted(items_data, key=lambda x: x['name'])

class ChatRequest(BaseModel):
    session_id: str
    message: str
    action: Optional[str] = None
    data: Optional[Dict[str, Any]] = None
    message_id: Optional[str] = None

# 全局变量存储会话数据
sessions: Dict[str, ChatSession] = {}
ai_generator = AIDocGenerator(api_key=os.environ.get("OPENROUTER_API_KEY", ""))

# --- Helper Functions ---
def get_or_create_session(session_id: str) -> ChatSession:
    if session_id not in sessions:
        logger.info(f"✨ Creating new session: {session_id}")
        sessions[session_id] = ChatSession(session_id)
    return sessions[session_id]

def get_session_upload_dir(session_id: str) -> Path:
    """Gets and creates the upload directory for a session."""
    upload_dir = Path("uploads") / session_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir

# --- API Endpoints ---
@app.get("/", response_class=HTMLResponse)
async def main_page(request: Request):
    """主页面"""
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/api/templates", response_model=List[Dict[str, str]])
async def get_templates():
    """获取可用的文档模板列表"""
    templates_list = []
    template_dir = Path("templates")
    template_dir.mkdir(exist_ok=True)

    # Add docx and doc files from the default 'templates' directory
    for ext in ["*.docx", "*.doc"]:
        for f in template_dir.glob(ext):
            templates_list.append({"name": f.name, "path": str(f)})
            
    if not templates_list:
        # To prevent errors on a clean setup, provide a placeholder
        return [{"name": "无可用模板", "path": ""}]
        
    return templates_list

@app.post("/api/generate")
async def generate_document_handler(
    session_id: str = Form(...),
    doc_id: str = Form(...),
    template_path: str = Form(...),
    json_data: str = Form(...),
    additional_docs: List[UploadFile] = File([])
):
    """
    Handles the primary document generation request with multimodal input.
    """
    session = get_or_create_session(session_id)
    doc_item = session.document_items.get(doc_id)
    if not doc_item:
        raise HTTPException(status_code=404, detail="Document item not found.")

    upload_dir = get_session_upload_dir(session_id)
    attachment_paths = []
    direct_json = None
    
    # --- Workflow Logic ---
    # Prioritize attachments over direct JSON input. This feels more intuitive for the user.
    # If files are uploaded, they are the source of truth for generation.

    # 1. Save any uploaded context files
    for file in additional_docs:
        file_path = upload_dir / f"context_{uuid.uuid4().hex[:8]}_{file.filename}"
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())
        attachment_paths.append(str(file_path))
        logger.info(f"💾 Saved context file for generation: {file_path}")

    # 2. Decide on the data source
    if not attachment_paths:
        # No attachments, so try to use the JSON data from the textarea
        logger.info("No attachments provided. Checking for direct JSON input.")
        if json_data and json_data.strip() and json_data.strip() != "{}":
            try:
                direct_json = json.loads(json_data)
                logger.info("✅ Using valid JSON data provided by user.")
            except json.JSONDecodeError:
                logger.error("⚠️ Invalid JSON provided by user and no attachments. Generation will fail.", exc_info=True)
                raise HTTPException(status_code=400, detail="Provided JSON is invalid and no context files were uploaded.")
        else:
            # No data source at all
            logger.error("❌ No data source provided. User sent neither files nor JSON.")
            raise HTTPException(status_code=400, detail="You must upload context documents or provide JSON data to generate a document.")
    else:
        # Attachments are present, they will be used as the primary source.
        logger.info(f"✅ Attachments found ({len(attachment_paths)} files). They will be used for AI data extraction.")


    output_dir = Path("generated_docs")
    output_dir.mkdir(exist_ok=True)
    # Sanitize the doc_item.name for the filename
    safe_name = "".join(c for c in doc_item.name if c.isalnum() or c in (' ', '_')).rstrip()
    output_path = output_dir / f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    try:
        success = ai_generator.run_generation(
            doc_template_path=template_path,
            output_path=str(output_path),
            attachment_paths=attachment_paths, # Will be empty if direct_json is used
            direct_json_data=direct_json      # Will be None if attachments are used
        )

        if success:
            doc_item.status = "已完成"
            doc_item.filled_document_path = str(output_path)
            response_message = f"✅ 文档 '{doc_item.name}' 已成功生成！"
        else:
            doc_item.status = "生成失败"
            doc_item.error_info = "AI生成过程出错或未能返回有效数据，请检查日志。"
            response_message = f"❌ 生成文档 '{doc_item.name}' 失败。"
    
    except Exception as e:
        logger.error(f"Error during document generation for '{doc_item.name}': {e}", exc_info=True)
        doc_item.status = "生成失败"
        doc_item.error_info = str(e)
        response_message = f"❌ 生成文档 '{doc_item.name}' 时发生严重错误。"
    
    finally:
        # Clean up temporary context files
        for path in attachment_paths:
            try:
                os.remove(path)
                logger.info(f"🗑️ Removed temp context file: {path}")
            except OSError as e:
                logger.error(f"Error removing temp file {path}: {e}")

    return JSONResponse({
        "message": response_message,
        "dashboard": session.get_dashboard_data()
    })

@app.post("/api/chat/message")
async def chat_handler(req: ChatRequest):
    """处理所有聊天、动作和文档生成请求"""
    session = get_or_create_session(req.session_id)
    response_message = ""
    response_options = []
    
    if req.action == "associate_template":
        data = req.data or {}
        doc_id = data.get("doc_id")
        template_path = data.get("template_path")
        doc_item = session.document_items.get(doc_id)

        if not doc_item or not template_path:
            raise HTTPException(status_code=400, detail="Document item and template path are required.")
        
        doc_item.matched_template_path = template_path
        response_message = f"✅ 模板已成功关联到项目 '{doc_item.name}'。"

    elif req.action == "reset_item":
        data = req.data or {}
        doc_id = data.get("doc_id")
        doc_item = session.document_items.get(doc_id)

        if not doc_item:
            raise HTTPException(status_code=404, detail="Document item not found.")
        
        # Reset the item's state
        doc_item.status = "待处理"
        doc_item.filled_document_path = None
        doc_item.error_info = None
        doc_item.matched_template_path = None
        
        logger.info(f"🔄 Item '{doc_item.name}' ({doc_id}) has been reset.")
        response_message = f"项目 '{doc_item.name}' 已重置。"

    elif req.message == '你好':
        response_message = "您好！我是AI文档生成助手。请从下方选择您要上传的内容类型，或直接在聊天框中向我提问。"
        response_options = [
            {"text": "项目竣工清单", "action": "upload_completion_list", "message_id": "init_1"},
            {"text": "多个模板", "action": "upload_templates", "message_id": "init_2"},
        ]
    else:
        response_message = f"我收到了您的消息: '{req.message}'. 目前通用聊天功能正在开发中。"

    return JSONResponse({
        "message": response_message,
        "options": response_options,
        "dashboard": session.get_dashboard_data()
    })

@app.post("/api/upload")
async def upload_file_handler(
    session_id: str = Form(...),
    upload_type: str = Form(...),
    doc_id: Optional[str] = Form(None),
    file: UploadFile = File(...)
):
    """文件上传处理"""
    session = get_or_create_session(session_id)
    
    # Create session-specific upload directory
    upload_dir = get_session_upload_dir(session_id)
    
    file_path = upload_dir / file.filename
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
        
    response_message = f"上传文件 '{file.filename}' 失败。"
    try:
        if upload_type == "upload_item_template":
            if not doc_id:
                raise HTTPException(status_code=400, detail="doc_id is required for this upload type")
            response_message = await process_item_template(session, str(file_path), file.filename, doc_id)
        elif upload_type == "upload_completion_list":
            response_message = await process_completion_list(session, str(file_path), file.filename)
        elif upload_type == "upload_templates":
            response_message = await process_templates(session, str(file_path), file.filename)
        elif upload_type == "upload_filled_doc":
            if not doc_id:
                 raise HTTPException(status_code=400, detail="doc_id is required for uploading a filled doc")
            
            output_dir = Path("generated_docs")
            output_dir.mkdir(exist_ok=True)
            # Use a unique name to avoid overwrites
            new_filename = f"{Path(file.filename).stem}_{doc_id}_{datetime.now().strftime('%Y%m%d%H%M%S')}{Path(file.filename).suffix}"
            saved_path = output_dir / new_filename
            
            # Since file is already saved in uploads, move it
            os.rename(file_path, saved_path)

            doc_item = session.document_items.get(doc_id)
            if doc_item:
                doc_item.status = "已完成"
                doc_item.matched_template_path = "Manually Uploaded" # Indicate manual override
                doc_item.filled_document_path = str(saved_path)
                response_message = f"✅ 已为 '{doc_item.name}' 上传并归档已填写的文档。"
            else:
                response_message = f"⚠️ 找不到ID为 {doc_id} 的项目，但文件已保存。"

        else:
            response_message = f"未知的上传类型: {upload_type}"
    except Exception as e:
        logger.error(f"Error processing upload for type '{upload_type}': {e}", exc_info=True)
        response_message = f"处理文件 '{file.filename}' 时发生错误。"

    return JSONResponse({
        "message": response_message,
        "dashboard": session.get_dashboard_data(),
        "options": []
    })

@app.get("/api/download/{session_id}/{doc_id}")
async def download_generated_file(session_id: str, doc_id: str):
    """下载已生成的文档"""
    session = get_or_create_session(session_id)
    doc_item = session.document_items.get(doc_id)

    if not doc_item or not doc_item.filled_document_path or not os.path.exists(doc_item.filled_document_path):
        raise HTTPException(status_code=404, detail="文件未找到或尚未生成。")

    return FileResponse(
        path=doc_item.filled_document_path,
        filename=os.path.basename(doc_item.filled_document_path),
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# --- File Processing Logic ---

async def process_item_template(session: ChatSession, file_path: str, filename: str, doc_id: str) -> str:
    """为特定项目关联模板"""
    if not (filename.endswith(".doc") or filename.endswith(".docx")):
        return f"❌ 模板文件格式不支持: '{filename}'。请上传 .doc 或 .docx 文件。"
    
    doc_item = session.document_items.get(doc_id)
    if not doc_item:
        return f"❌ 未找到项目ID: {doc_id}"

    # 将模板保存到共享的'templates'目录中，以确保其可用
    template_dir = Path("templates")
    template_dir.mkdir(exist_ok=True)
    
    # 为避免命名冲突，可以在文件名中加入doc_id
    new_filename = f"{Path(filename).stem}_{doc_id}{Path(filename).suffix}"
    target_path = template_dir / new_filename
    
    import shutil
    shutil.copy(file_path, target_path)

    # 更新文档项目
    doc_item.matched_template_path = str(target_path)
    
    # 将其也添加到会话中的通用模板列表中
    session.templates[new_filename] = str(target_path)

    return f"✅ 已为项目 '{doc_item.name}' 成功关联模板 '{filename}'。"

async def process_completion_list(session: ChatSession, file_path: str, filename: str) -> str:
    """从清单文件中提取项目并更新仪表板"""
    items = []
    try:
        if filename.endswith('.docx'):
            doc = Document(file_path)
            # Try to extract from tables first
            for table in doc.tables:
                for row in table.rows:
                    cell_text = " ".join(cell.text.strip() for cell in row.cells).strip()
                    if cell_text:
                        items.append(cell_text)
            # If no tables, extract from paragraphs
            if not items:
                for para in doc.paragraphs:
                    if para.text.strip():
                        items.append(para.text.strip())
        elif filename.endswith('.json'):
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Expecting a list of strings or a dict with an "items" key
                items = data if isinstance(data, list) else data.get("items", [])
        elif filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                items = [line.strip() for line in f if line.strip()]
        
        # Filter out duplicates and very short lines
        unique_items = sorted(list(set([item for item in items if len(item) > 2])))
        
        if not unique_items:
            return f"在 '{filename}' 中未找到有效的项目。"

        # Update session dashboard
        for item_name in unique_items:
            # Avoid adding duplicate names
            if not any(item.name == item_name for item in session.document_items.values()):
                doc_id = f"doc_{uuid.uuid4().hex[:8]}"
                session.document_items[doc_id] = DocumentItem(name=item_name, doc_id=doc_id)

        return f"✅ 成功处理清单 '{filename}'，找到并更新了 {len(unique_items)} 个待办项目。"
    except Exception as e:
        logger.error(f"Error processing completion list '{filename}': {e}", exc_info=True)
        return f"❌ 处理清单 '{filename}' 时出错。"

async def process_templates(session: ChatSession, file_path: str, filename: str) -> str:
    """处理上传的模板文件"""
    if not (filename.endswith(".doc") or filename.endswith(".docx")):
        return f"❌ 模板文件格式不支持: '{filename}'。请上传 .doc 或 .docx 文件。"
        
    # Store template path in session for later use
    session.templates[filename] = file_path
    
    # Also copy to the main 'templates' directory to make it available for all sessions
    # This is a design choice - templates are shared.
    default_template_dir = Path("templates")
    default_template_dir.mkdir(exist_ok=True)
    
    target_path = default_template_dir / filename
    if not target_path.exists():
        import shutil
        shutil.copy(file_path, target_path)
        return f"✅ 模板 '{filename}' 已成功上传并可供使用。"
    else:
        return f"ℹ️ 模板 '{filename}' 已存在，无需重复上传。"

# --- Main Execution ---
if __name__ == '__main__':
    logger.info("🚀 启动AI文档生成器Web应用...")
    # It's recommended to run with Uvicorn directly for more control
    # Example: uvicorn app:app --host 0.0.0.0 --port 8000 --reload
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True) 