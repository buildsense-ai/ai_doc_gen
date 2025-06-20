#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
主程序：AI文档生成器
支持三阶段流程：DOC转换 → 模板分析 → JSON输入 → 文档生成
"""

import os
import json
import logging
import subprocess
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from openai import OpenAI
import base64
import mimetypes
import fitz  # PyMuPDF
from docx import Document as DocxDocument

# Load environment variables from .env file if it exists
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # python-dotenv not installed, skip .env file loading
    pass

# Import prompts
from prompt_utils import get_fill_data_prompt, get_multimodal_extraction_prompt

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)
logger = logging.getLogger(__name__)

# Define a directory for uploads and temporary files
UPLOADS_DIR = "uploads"
if not os.path.exists(UPLOADS_DIR):
    os.makedirs(UPLOADS_DIR)

class AIDocGenerator:
    """AI文档生成器 - 支持DOC转换"""
    
    def __init__(self, api_key: str):
        """初始化OpenRouter客户端"""
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self.model = "google/gemini-2.5-pro-preview"
        logger.info("🤖 AI生成器初始化完成")
    
    def _extract_json_from_response(self, response_content: str) -> str:
        """
        Extract JSON string from AI response content.
        Handles various formats like markdown code blocks, plain JSON, etc.
        """
        if not response_content or not response_content.strip():
            raise ValueError("AI response content is empty")
        
        content = response_content.strip()
        
        # Try to extract from markdown code block
        if "```json" in content:
            try:
                start = content.find("```json") + 7
                end = content.find("```", start)
                if end != -1:
                    json_str = content[start:end].strip()
                    if json_str:
                        return json_str
            except Exception:
                pass
        
        # Try to extract from single backticks
        if content.startswith("`") and content.endswith("`"):
            json_str = content.strip("`").strip()
            if json_str:
                return json_str
        
        # Try to find JSON object boundaries
        start_idx = content.find("{")
        if start_idx != -1:
            # Find the matching closing brace
            brace_count = 0
            for i, char in enumerate(content[start_idx:], start_idx):
                if char == "{":
                    brace_count += 1
                elif char == "}":
                    brace_count -= 1
                    if brace_count == 0:
                        json_str = content[start_idx:i+1]
                        # Validate it's proper JSON
                        try:
                            json.loads(json_str)
                            return json_str
                        except json.JSONDecodeError:
                            continue
        
        # If all else fails, try the content as-is
        try:
            json.loads(content)
            return content
        except json.JSONDecodeError:
            raise ValueError(f"Could not extract valid JSON from AI response: {content[:200]}...")

    def convert_doc_to_docx(self, doc_path: str) -> str:
        """
        使用LibreOffice将.doc文件转换为.docx文件
        
        Args:
            doc_path: .doc文件路径
            
        Returns:
            转换后的.docx文件路径
        """
        logger.info("🔄 开始DOC到DOCX转换...")
        
        if not os.path.exists(doc_path):
            logger.error(f"❌ DOC文件不存在: {doc_path}")
            raise FileNotFoundError(f"DOC文件不存在: {doc_path}")
        
        # 生成输出文件名
        docx_path = doc_path.replace('.doc', '_converted.docx')
        
        try:
            # 检查LibreOffice是否可用
            logger.info("🔍 检查LibreOffice可用性...")
            
            # 尝试多个可能的LibreOffice路径
            libreoffice_paths = [
                '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS
                'libreoffice',  # Linux/Windows PATH
                'soffice',  # 备用命令
            ]
            
            libreoffice_cmd = None
            for path in libreoffice_paths:
                try:
                    result = subprocess.run([path, '--version'], 
                                          capture_output=True, 
                                          text=True, 
                                          timeout=10)
                    if result.returncode == 0:
                        libreoffice_cmd = path
                        logger.info(f"✅ 找到LibreOffice: {path}")
                        break
                except (FileNotFoundError, subprocess.TimeoutExpired):
                    continue
            
            if not libreoffice_cmd:
                logger.error("❌ 未找到LibreOffice，请确保已安装LibreOffice")
                raise RuntimeError("LibreOffice未安装或不可用")
            
            # 执行转换
            logger.info(f"📄 正在转换: {doc_path} -> {docx_path}")
            
            # 删除已存在的输出文件
            if os.path.exists(docx_path):
                os.remove(docx_path)
                logger.info("🗑️ 删除已存在的转换文件")
            
            # LibreOffice转换命令
            cmd = [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', os.path.dirname(doc_path),
                doc_path
            ]
            
            logger.info(f"🔧 执行命令: {' '.join(cmd)}")
            
            result = subprocess.run(cmd, 
                                  capture_output=True, 
                                  text=True, 
                                  timeout=30)
            
            if result.returncode != 0:
                logger.error(f"❌ LibreOffice转换失败: {result.stderr}")
                raise RuntimeError(f"LibreOffice转换失败: {result.stderr}")
            
            # 检查转换后的文件
            expected_docx = doc_path.replace('.doc', '.docx')
            if os.path.exists(expected_docx):
                # 重命名为我们期望的文件名
                if expected_docx != docx_path:
                    os.rename(expected_docx, docx_path)
                
                logger.info(f"✅ 转换成功: {docx_path}")
                return docx_path
            else:
                logger.error(f"❌ 转换后的文件未找到: {expected_docx}")
                raise RuntimeError("转换后的文件未找到")
                
        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice转换超时")
            raise RuntimeError("LibreOffice转换超时")
        except Exception as e:
            logger.error(f"❌ 转换过程中出错: {e}")
            raise
    
    def stage1_analyze_template(self, template_path: str) -> Dict[str, str]:
        """
        阶段1：确定性地分析Word模板，提取带有位置信息的结构。
        
        Args:
            template_path: .docx模板文件路径

        Returns:
            一个字典，其中键是单元格的唯一标识符，值是单元格的文本内容。
        """
        logger.info("🔍 阶段1：开始确定性模板结构分析...")
        
        try:
            doc = Document(template_path)
            template_structure = {}
            
            logger.info(f"📄 正在读取模板文件: {template_path}")
            
            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        cell_key = f"table_{i}_row_{j}_col_{k}"
                        template_structure[cell_key] = cell.text.strip()
            
            logger.info(f"✅ 成功提取 {len(template_structure)} 个单元格的结构信息。")
            # Log a snippet of the extracted structure for verification
            structure_snippet = json.dumps(dict(list(template_structure.items())[:5]), ensure_ascii=False, indent=2)
            logger.info(f"  结构实例:\n{structure_snippet}")

            return template_structure
            
        except Exception as e:
            logger.error(f"❌ 阶段1错误: {e}")
            raise

    def stage2_load_json_data(self, json_file_path: str) -> Dict[str, Any]:
        """
        阶段2：从JSON文件加载数据
        """
        logger.info("📂 阶段2：开始加载JSON数据...")
        
        try:
            if not os.path.exists(json_file_path):
                logger.error(f"❌ JSON文件不存在: {json_file_path}")
                raise FileNotFoundError(f"JSON文件不存在: {json_file_path}")
            
            logger.info(f"📄 正在读取JSON文件: {json_file_path}")
            
            with open(json_file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            logger.info(f"✅ 成功加载 {len(data)} 个数据字段。")
            for key, value in data.items():
                preview = str(value)[:70] + "..." if len(str(value)) > 70 else str(value)
                logger.info(f"   📌 {key}: {preview}")
            
            return data
            
        except json.JSONDecodeError as e:
            logger.error(f"❌ 阶段2错误: JSON文件格式无效 - {e}")
            raise
        except Exception as e:
            logger.error(f"❌ 阶段2错误: {e}")
            raise
    
    def stage2_1_ai_extract_data_from_sources(self, attachment_paths: List[str]) -> Dict[str, Any]:
        """
        Stage 2.1: Use multimodal AI to extract data from various documents and images.
        """
        logger.info("🧠 Stage 2.1: Kicking off multimodal AI data extraction...")
        
        try:
            # This is a sample schema. In a real app, this might come from the template
            # or a user configuration. For now, we'll use a schema based on sample_input.json
            fields_to_extract = json.dumps({
                "serial_number": "示例: GZ-FH-2025-001",
                "project_name": "示例: 历史建筑修复项目",
                "review_date": "示例: 2025-01-25",
                "original_condition_review": "建筑物原始状态的描述。",
                "damage_assessment_review": "发现的任何损伤的详细评估。",
                "repair_plan_review": "拟定的修复计划。",
                "project_lead": "项目负责人姓名。",
                "reviewer": "审核人员姓名。",
                "damage_photos_path": "损伤照片文件路径列表，如果有的话。",
                "site_photos_path": "现场照片文件路径列表，如果有的话。",
                "attachments": "相关图像文件路径列表，如果有的话。为每个图像提供描述性标题。"
            }, indent=2, ensure_ascii=False)

            prompt = get_multimodal_extraction_prompt(fields_to_extract)

            # Build the message with text and images
            content_parts = [{"type": "text", "text": prompt}]
            
            # --- Unified File Processing Loop ---
            image_paths_for_prompt = []
            temp_text_files = []

            for file_path in attachment_paths:
                file_name = os.path.basename(file_path)
                logger.info(f"📄 Processing attachment: {file_name}")

                try:
                    if file_path.endswith(('.txt', '.md', '.json')):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            file_content = f.read()
                        text_part = f"\n\n--- Content from {file_name} ---\n{file_content}\n--- End of Content ---"
                        content_parts[0]["text"] += text_part

                    elif file_path.endswith('.docx'):
                        doc = DocxDocument(file_path)
                        full_text = "\n".join([p.text for p in doc.paragraphs])
                        text_part = f"\n\n--- Content from {file_name} ---\n{full_text}\n--- End of Content ---"
                        content_parts[0]["text"] += text_part

                    elif file_path.endswith('.pdf'):
                        doc = fitz.open(file_path)
                        full_text = ""
                        for page_num, page in enumerate(doc):
                            full_text += page.get_text()
                            # Extract images from PDF
                            img_list = page.get_images(full=True)
                            for img_index, img in enumerate(img_list):
                                xref = img[0]
                                base_image = doc.extract_image(xref)
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]
                                
                                # Save image to a temporary file
                                temp_image_filename = f"pdf_{os.path.splitext(file_name)[0]}_p{page_num+1}_img{img_index}.{image_ext}"
                                temp_image_path = os.path.join(UPLOADS_DIR, temp_image_filename)
                                with open(temp_image_path, "wb") as f:
                                    f.write(image_bytes)
                                
                                image_paths_for_prompt.append(temp_image_path)
                                logger.info(f"🖼️  Extracted image from PDF: {temp_image_path}")
                        
                        text_part = f"\n\n--- Content from {file_name} ---\n{full_text}\n--- End of Content ---"
                        content_parts[0]["text"] += text_part
                        doc.close()

                    else: # Assumes it's an image if not a text-based file
                        mime_type, _ = mimetypes.guess_type(file_path)
                        if mime_type and mime_type.startswith('image/'):
                            image_paths_for_prompt.append(file_path)
                        else:
                            logger.warning(f"⚠️ Unsupported file type, skipping: {file_name}")

                except Exception as e:
                    logger.error(f"❌ Error processing file {file_path}: {e}", exc_info=True)


            # Add all collected images to the prompt
            for image_path in image_paths_for_prompt:
                try:
                    mime_type, _ = mimetypes.guess_type(image_path)
                    with open(image_path, "rb") as image_file:
                        base64_image = base64.b64encode(image_file.read()).decode('utf-8')
                    
                    image_url = f"data:{mime_type};base64,{base64_image}"
                    
                    # Add a reference in the text part with Chinese description
                    content_parts[0]["text"] += f"\n\n--- 附加图像 (文件路径: {image_path}) ---"
                    
                    content_parts.append({
                        "type": "image_url",
                        "image_url": {"url": image_url}
                    })
                    logger.info(f"🖼️  Added image {image_path} to AI prompt.")
                except Exception as e:
                    logger.warning(f"⚠️ Could not process image file {image_path}: {e}")

            logger.info("🧠 Calling multimodal AI to extract structured data... (This may take a moment)")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": content_parts}],
                temperature=0.1
            )
            
            # Clean up extracted text files
            for path in temp_text_files:
                try:
                    os.remove(path)
                except OSError as e:
                    logger.error(f"Error removing temp text file {path}: {e}")
            
            # Extract and parse the JSON from the AI's response
            if response.choices[0].message.content:
                json_string = self._extract_json_from_response(response.choices[0].message.content)
                extracted_data = json.loads(json_string)
                
                logger.info(f"✅ AI successfully extracted data. Keys: {list(extracted_data.keys())}")
                return extracted_data
            else:
                raise ValueError("AI returned an empty response.")
                
        except Exception as e:
            logger.error(f"❌ Stage 2.1 Error: {e}", exc_info=True)
            raise

    def stage2_5_ai_generate_fill_data(self, structured_template: Dict[str, str], input_data: Dict[str, Any]) -> Dict[str, str]:
        """
        阶段2.5：使用AI将输入数据智能映射到模板结构，生成用于填充的最终数据。
        
        Args:
            structured_template: 从阶段1获得的模板结构
            input_data: 从阶段2获得的输入数据
            
        Returns:
            一个字典，键是单元格的唯一标识符，值是待填充的数据。
        """
        logger.info("🧠 阶段2.5：开始AI字段映射和数据生成...")
        
        try:
            # 构建AI映射提示
            prompt = get_fill_data_prompt(
                json.dumps(structured_template, ensure_ascii=False, indent=2),
                json.dumps(input_data, ensure_ascii=False, indent=2)
            )
            
            logger.info("🧠 正在调用AI生成填充数据... (这可能需要一些时间)")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                extra_headers={
                    "HTTP-Referer": "ai-doc-generator",
                    "X-Title": "AI Document Generator",
                },
                temperature=0.1, # Use low temperature for more predictable output
            )

            # Log the full response for debugging purposes
            logger.info(f"🔍 Raw AI Response (for debugging):\n{response.model_dump_json(indent=2)}")

            # Gracefully handle invalid or empty responses from the API
            if not response or not response.choices:
                logger.error("❌ AI响应无效或不包含 'choices' 字段。")
                logger.error(f"   Full API Response: {response}")
                logger.warning("⚠️ AI字段映射失败，将返回空数据。")
                return {}

            message = response.choices[0].message
            if not message or not message.content:
                logger.error("❌ AI响应的消息内容为空。")
                logger.error(f"   Full choice object: {response.choices[0].model_dump_json(indent=2)}")
                logger.warning("⚠️ AI字段映射失败，将返回空数据。")
                return {}
            
            # 解析返回的JSON
            json_text = message.content
            if "```json" in json_text:
                json_text = json_text.split("```json")[1].split("```")[0]
            elif json_text.startswith("`") and json_text.endswith("`"):
                json_text = json_text.strip("`")

            fill_data = json.loads(json_text.strip())
            
            # Check for attachments in the AI response
            if '__attachments__' in fill_data:
                logger.info(f"🎯 AI生成了 {len(fill_data['__attachments__'])} 个附件引用")
                for i, att in enumerate(fill_data['__attachments__']):
                    logger.info(f"   📎 附件 {i+1}: {att}")
            else:
                logger.info("ℹ️ AI响应中未包含附件数据")
            
            logger.info(f"✅ AI成功生成 {len(fill_data)} 个字段的映射:")
            for key, value in fill_data.items():
                if key == '__attachments__':
                    logger.info(f"   🔗 {key} -> [包含 {len(value)} 个附件]")
                else:
                    preview = str(value)[:70] + "..." if len(str(value)) > 70 else str(value)
                    logger.info(f"   🔗 {key} -> '{preview}'")
            
            return fill_data
            
        except json.JSONDecodeError as e:
            logger.error(f"❌ 阶段2.5错误: AI返回的JSON无效 - {e}")
            logger.error(f"   Raw AI Response: {json_text}")
            logger.warning("⚠️ AI字段映射失败，将返回空数据。")
            return {}
        except Exception as e:
            logger.error(f"❌ 阶段2.5错误: {e}")
            logger.warning("⚠️ AI字段映射失败，将返回空数据。")
            return {}

    def stage3_fill_template(self, template_path: str, output_path: str, fill_data: Dict[str, str]):
        """
        阶段3：根据AI生成的填充数据，确定性地填充模板。
        
        Args:
            template_path: .docx模板文件路径
            output_path: 输出文件路径
            fill_data: 从阶段2.5获得的填充数据
        """
        logger.info("📝 阶段3：开始确定性模板填充...")
        
        if not os.path.exists(template_path):
            logger.error(f"❌ 模板文件未找到: {template_path}")
            raise FileNotFoundError(f"模板文件未找到: {template_path}")

        try:
            doc = Document(template_path)
            filled_fields_count = 0
            
            # Extract attachments before processing other fields
            attachments_data = fill_data.pop('__attachments__', [])
            logger.info(f"📎 发现 {len(attachments_data)} 个附件待处理")
            if attachments_data:
                for i, att in enumerate(attachments_data):
                    logger.info(f"   附件 {i+1}: {att.get('title', 'N/A')} -> {att.get('path', 'N/A')}")
            
            # 创建一份待填充字段的副本，用于追踪
            remaining_to_fill = set(fill_data.keys())

            for i, table in enumerate(doc.tables):
                for j, row in enumerate(table.rows):
                    for k, cell in enumerate(row.cells):
                        cell_key = f"table_{i}_row_{j}_col_{k}"
                        if cell_key in fill_data:
                            fill_value = str(fill_data[cell_key])
                            # 清空单元格原有内容（如占位符），然后填充
                            cell.text = fill_value
                            logger.info(f"   ✏️ 填充 {cell_key}: '{fill_value[:50]}...'")
                            filled_fields_count += 1
                            remaining_to_fill.discard(cell_key)

            # Add attachments at the end of the document
            if attachments_data:
                logger.info(f"📎 开始附加 {len(attachments_data)} 个文件到文档末尾...")
                # Add a page break before attachments if document is not empty
                if len(doc.paragraphs) > 0 or len(doc.tables) > 0:
                    doc.add_page_break()
                
                # Add a main heading for attachments section  
                # Use paragraph instead of heading to avoid style issues
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("附件")
                run.bold = True
                run.font.size = Pt(16)  # Larger font size like a heading
                    
                for i, attachment in enumerate(attachments_data, 1):
                    title = attachment.get('title', f'附件 {i}')
                    path = attachment.get('path')
                    
                    if path and os.path.exists(path):
                        try:
                            # Add numbered attachment heading using paragraph  
                            heading_para = doc.add_paragraph()
                            heading_run = heading_para.add_run(f"{i}. {title}")
                            heading_run.bold = True
                            heading_run.font.size = Pt(14)  # Slightly smaller than main heading
                            
                            # Determine optimal image size based on file size and type
                            mime_type, _ = mimetypes.guess_type(path)
                            if mime_type and mime_type.startswith('image/'):
                                # Add the image with reasonable sizing
                                doc.add_picture(path, width=Inches(6.0))
                                logger.info(f"   ✅ 已附加图片: {path}")
                            else:
                                # For non-image files, add a note
                                p = doc.add_paragraph(f"文件: {os.path.basename(path)}")
                                logger.info(f"   📄 已添加文件引用: {path}")
                                
                        except Exception as e:
                            logger.error(f"   ❌ 附加文件失败 {path}: {e}")
                            # Add error note in document
                            doc.add_paragraph(f"⚠️ 无法显示附件: {os.path.basename(path) if path else 'Unknown'}")
                    else:
                        logger.warning(f"   ⚠️ 附件文件未找到或路径无效: {path}")
                        # Add missing file note in document
                        doc.add_paragraph(f"⚠️ 附件文件未找到: {os.path.basename(path) if path else 'Unknown'}")

            # 保存文档
            doc.save(output_path)
            
            logger.info(f"✅ 文档已成功生成: {output_path}")
            logger.info(f"📊 共填充 {filled_fields_count} / {len(fill_data)} 个AI映射的字段。")

            # 检查是否有任何映射的字段未被填充
            if remaining_to_fill:
                logger.warning("⚠️ 以下由AI映射的字段在模板中未找到对应的单元格并被跳过：")
                for key in remaining_to_fill:
                    logger.warning(f"   - {key}")
            else:
                logger.info("✅ 所有AI映射的字段都已成功填充。")
            
        except Exception as e:
            logger.error(f"❌ 阶段3错误: {e}")
            raise

    def run_generation(
        self, 
        doc_template_path: str, 
        output_path: str, 
        attachment_paths: Optional[List[str]] = None,
        direct_json_data: Optional[Dict[str, Any]] = None
    ):
        """
        Runs the full document generation process.

        Supports two workflows:
        1. If 'direct_json_data' is provided, it uses that data directly.
        2. If 'direct_json_data' is None, it runs AI extraction on 'attachment_paths'.
        """
        logger.info("🚀 Starting unified document generation process...")
        
        try:
            # Stage 0: Convert .doc to .docx if necessary
            if doc_template_path.lower().endswith('.doc'):
                logger.info(f"📄 Detected .doc template. Attempting conversion for: {doc_template_path}")
                processed_template_path = self.convert_doc_to_docx(doc_template_path)
            else:
                processed_template_path = doc_template_path

            # Stage 1 is always required to know the template structure
            template_structure = self.stage1_analyze_template(processed_template_path)
            
            input_data = {}
            if direct_json_data:
                logger.info("📄 Using user-provided JSON data directly.")
                input_data = direct_json_data
            elif attachment_paths:
                logger.info("🧠 No direct JSON provided, starting AI extraction from attachments.")
                # Stage 2.1: Use AI to extract data from attachments
                input_data = self.stage2_1_ai_extract_data_from_sources(
                    attachment_paths=attachment_paths
                )
            else:
                raise ValueError("Generation failed: You must provide either direct JSON data or attachment files.")

            # Stage 2.5: Use AI to map extracted/provided data to the template structure
            fill_data = self.stage2_5_ai_generate_fill_data(
                structured_template=template_structure,
                input_data=input_data
            )
            
            # Stage 3: Fill the Word template with the final data
            self.stage3_fill_template(
                template_path=processed_template_path,
                output_path=output_path,
                fill_data=fill_data
            )
            
            logger.info(f"✅ Document generation complete: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Document generation failed: {e}", exc_info=True)
            return False

    def run_complete_workflow(self, doc_template_path: str, json_input_path: str, output_path: str):
        """
        运行完整的3阶段工作流（从模板和JSON文件）
        """
        logger.info("🚀 开始完整的AI文档生成流程")
        logger.info("=" * 60)
        
        start_time = datetime.now()

        # Create a dedicated directory for this generation job's intermediate files
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        job_dir = os.path.join("generated_docs", f"job_{timestamp}")
        os.makedirs(job_dir, exist_ok=True)
        logger.info(f"📁 Created job directory: {job_dir}")

        # 基于输入模板名称，创建中间文件的路径
        base_name = os.path.splitext(os.path.basename(doc_template_path))[0]
        structure_output_path = os.path.join(job_dir, f"{base_name}_template_structure.json")
        fill_data_output_path = os.path.join(job_dir, f"{base_name}_filled_data.json")
        
        try:
            # 阶段 0：DOC转DOCX (如果需要)
            if doc_template_path.endswith('.doc'):
                docx_template_path = self.convert_doc_to_docx(doc_template_path)
            else:
                docx_template_path = doc_template_path
            logger.info("=" * 30)
            
            # 阶段 1：分析模板结构
            structured_template = self.stage1_analyze_template(docx_template_path)
            
            # 保存中间结果1: 模板结构JSON
            try:
                with open(structure_output_path, 'w', encoding='utf-8') as f:
                    json.dump(structured_template, f, ensure_ascii=False, indent=4)
                logger.info(f"💾 中间结果已保存: {structure_output_path}")
            except Exception as e:
                logger.error(f"❌ 保存模板结构JSON时出错: {e}")

            logger.info("=" * 30)
            
            # 阶段 2：加载JSON数据
            input_data = self.stage2_load_json_data(json_input_path)
            logger.info("=" * 30)
            
            # 阶段 2.5：AI生成填充数据
            fill_data = self.stage2_5_ai_generate_fill_data(structured_template, input_data)

            # 保存中间结果2: AI生成的填充数据JSON
            if fill_data:
                try:
                    with open(fill_data_output_path, 'w', encoding='utf-8') as f:
                        json.dump(fill_data, f, ensure_ascii=False, indent=4)
                    logger.info(f"💾 中间结果已保存: {fill_data_output_path}")
                except Exception as e:
                    logger.error(f"❌ 保存填充数据JSON时出错: {e}")
            else:
                logger.warning("⚠️ AI未生成填充数据，跳过保存中间JSON文件。")
            
            logger.info("=" * 30)
            
            # 阶段 3：填充模板
            self.stage3_fill_template(docx_template_path, output_path, fill_data)
            
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            logger.info("=" * 60)
            logger.info("🎉 AI文档生成流程完成!")
            logger.info(f"⏱️ 总用时: {duration:.2f} 秒")
            logger.info(f"📄 输出文件: {output_path}")
            if docx_template_path != doc_template_path:
                logger.info(f"🔄 中间转换文件: {docx_template_path}")
            logger.info(f"📊 中间结构文件: {structure_output_path}")
            if fill_data:
                logger.info(f"🧠 中间填充数据: {fill_data_output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"❌ 工作流程中发生致命错误: {e}", exc_info=True)
            return False


def main():
    """主函数"""
    print("🚀 AI文档生成器 - 主程序")
    print("=" * 50)
    
    # --- 配置 ---
    # API Key从环境变量读取，确保安全性
    # 你可以从这里获取API Key: https://openrouter.ai/keys
    API_KEY = os.environ.get("OPENROUTER_API_KEY")
    
    if not API_KEY:
        logger.error("❌ 错误: 未找到 OPENROUTER_API_KEY 环境变量")
        logger.error("请设置环境变量:")
        logger.error("  macOS/Linux: export OPENROUTER_API_KEY='your-api-key-here'")
        logger.error("  Windows: set OPENROUTER_API_KEY=your-api-key-here")
        logger.error("或者创建 .env 文件并添加: OPENROUTER_API_KEY=your-api-key-here")
        return

    # 文件路径
    doc_template_path = "template_test2.doc"  # 使用.doc或.docx文件
    json_input_path = "sample_input2.json"
    output_path = f"AI生成文档_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # 检查输入文件
    if not os.path.exists(doc_template_path):
        logger.error(f"❌ 模板文件不存在: {doc_template_path}")
        return
    
    if not os.path.exists(json_input_path):
        logger.error(f"❌ JSON输入文件不存在: {json_input_path}")
        return
    
    # 初始化并运行工作流程
    try:
        generator = AIDocGenerator(API_KEY)
        success = generator.run_complete_workflow(
            doc_template_path=doc_template_path,
            json_input_path=json_input_path,
            output_path=output_path
        )
        
        if success:
            print(f"\n✅ 成功！生成的文档已保存至: {output_path}")
        else:
            print("\n❌ 失败！请检查上面的日志信息以了解详情。")

    except Exception as e:
        logger.error(f"❌ 初始化或运行生成器时发生未处理的错误: {e}", exc_info=True)


if __name__ == "__main__":
    # 检查是否要启动Web界面
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--web":
        # 启动Web界面
        import subprocess
        subprocess.run([sys.executable, "app.py"])
    elif len(sys.argv) > 1 and sys.argv[1] == "--cli":
        # 启动命令行界面
        main()
    else:
        # 默认启动Web界面
        print("🌐 启动Web界面...")
        print("如需使用命令行版本，请运行: python main.py --cli")
        print("=" * 50)
        import subprocess
        subprocess.run([sys.executable, "app.py"]) 